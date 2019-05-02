#!/usr/bin/env python
# coding: utf-8

# In[30]:


from selenium import webdriver
from selenium.webdriver.support.select import Select
import requests
from bs4 import BeautifulSoup
import docx

doc = docx.Document()
driver = webdriver.Chrome('C:/Users/user/chromedriver.exe')    #chromedriver location
driver.get('http://asbc.iis.sinica.edu.tw/index_range.htm')
soup=BeautifulSoup(driver.page_source)
driver.find_element_by_xpath("//body/form[@name='form1']//input[@name='Submit']").click()

# 查詢關鍵字網址: http://asbc.iis.sinica.edu.tw/index2.asp


# In[31]:


target=str(input("請輸入要找尋的文字："))
id=driver.find_element_by_name('inputword')
id.send_keys(target)
driver.find_element_by_xpath("//body/div/table//form[@name='form1']//input[@value='送出']").click()

string_len=len(target)


# In[32]:


# 要分開，要讓driver讀去page source
soup=BeautifulSoup(driver.page_source)


# In[33]:


# 顯示總共有多少筆

print(soup.find_all("p")[-1].text)

page=soup.find_all("p")[-1].text.replace("共有 ","").replace(" 筆資料","")
print("共有",int(page)//100+1," 頁")
sum=int(page)//100+1


# In[36]:


if string_len==1:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="501")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="501")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)
        
elif string_len==2:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="490")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="490")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)

else:
    print()

     # 第 1 頁 


# In[35]:


driver.find_element_by_xpath("//body/div[2]/a[10]").click()
if string_len==1:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="501")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="501")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)
   
elif string_len==2:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="490")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="490")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)

else:
    print()
    
    # 第 2 頁
    


# In[148]:


for i in range(sum-5):
    driver.find_element_by_xpath("//body/div[2]/a[11]").click()
    if string_len==1:
        soup=BeautifulSoup(driver.page_source)
        left=soup.find_all("div",align="right")            
        main=soup.find_all("b")                             # 主要搜尋字詞
        right=soup.find_all("td", width="501")
        list_len=len(left)
        for i in range(list_len):
            left=soup.find_all("div",align="right")
            main=soup.find_all("b")
            right=soup.find_all("td", width="501")
            paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
            doc.add_paragraph(paragraph)
            
    elif string_len==2:
        soup=BeautifulSoup(driver.page_source)
        left=soup.find_all("div",align="right")           
        main=soup.find_all("b")                             # 主要搜尋字詞
        right=soup.find_all("td", width="490")
        list_len=len(left)
        for i in range(list_len):
            left=soup.find_all("div",align="right")
            main=soup.find_all("b")
            right=soup.find_all("td", width="490")
            paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
            doc.add_paragraph(paragraph)
    else:
        print()
        
        # 中間頁數
        


# In[149]:


driver.find_element_by_xpath("//body/div[2]/a[10]").click()
if string_len==1:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="501")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="501")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)
   
elif string_len==2:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="490")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="490")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)

else:
    print()
        # 倒數第 3頁
        


# In[150]:


driver.find_element_by_xpath("//body/div[2]/a[9]").click()
if string_len==1:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="501")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="501")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)
   
elif string_len==2:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="490")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="490")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)

else:
    print()
        
        # 倒數第 2頁
        


# In[151]:


driver.find_element_by_xpath("//body/div[2]/a[8]").click()
if string_len==1:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="501")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="501")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)
   
elif string_len==2:
    soup=BeautifulSoup(driver.page_source)
    left=soup.find_all("div",align="right")           
    main=soup.find_all("b")                             # 主要搜尋字詞
    right=soup.find_all("td", width="490")
    list_len=len(left)
    for i in range(list_len):
        left=soup.find_all("div",align="right")
        main=soup.find_all("b")
        right=soup.find_all("td", width="490")
        paragraph=left[i].text.replace('more',''),main[0].text, right[i].text.replace('more','')
        doc.add_paragraph(paragraph)

else:
    print()
        # 最後一頁
        


# In[153]:


doc.save(target+'.docx')

